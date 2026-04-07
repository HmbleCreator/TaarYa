from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT = Path(r"C:\Users\amiku\Downloads\8th_Sem\TaarYa")
OUT_DIR = ROOT / "report_output"
OUTPUT = OUT_DIR / "TaarYa_Major_Project_Report_ES452.docx"
DTC_LOGO = ROOT / "docx_inspect" / "image1.png"
TAARYA_LOGO = ROOT / "static" / "TaarYaLogo.png"

DEPARTMENT = "Department of Computer Science & Engineering"


def set_cell_text(cell, text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction, placeholder=""):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = placeholder
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def set_page_number_format(section, fmt, start=None):
    sect_pr = section._sectPr
    for existing in sect_pr.xpath("./w:pgNumType"):
        sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def configure_section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.left_margin = Inches(1.5)


def configure_footer(section, fmt, start=None):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)

    run = paragraph.add_run(DEPARTMENT)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    paragraph.add_run("\t")
    add_field(paragraph, "PAGE", "1")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)
    set_page_number_format(section, fmt, start=start)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.line_spacing = 1.5
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.line_spacing = 1.5
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Times New Roman"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading3.font.size = Pt(12)
    heading3.font.italic = True
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.line_spacing = 1.5
    heading3.paragraph_format.space_before = Pt(4)
    heading3.paragraph_format.space_after = Pt(2)

    if "CodeBlock" not in doc.styles:
        code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(10)
        code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code_style.paragraph_format.left_indent = Inches(0.3)
        code_style.paragraph_format.right_indent = Inches(0.3)
        code_style.paragraph_format.space_after = Pt(3)


def add_centered(doc, text, size=12, bold=False, italic=False, spacing_after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.add_run("- " + text)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    return paragraph


def add_placeholder(doc, label, height_inches=2.4):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    row = table.rows[0]
    row.height = Inches(height_inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Inches(5.8)
    set_cell_text(
        cell,
        f"Space reserved for {label}.\nInsert the final diagram/image here.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        size=12,
    )
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row_cells[idx], str(value), align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    for line in code.strip().splitlines():
        paragraph = doc.add_paragraph(style="CodeBlock")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(line.rstrip())
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_cover_page(doc):
    if DTC_LOGO.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(DTC_LOGO), width=Inches(1.4))

    add_centered(doc, "TAARYA", size=18, bold=True, spacing_after=6)
    add_centered(
        doc,
        "An Intelligent Agentic RAG-Driven Architecture for Astronomical Star Catalogs",
        size=24,
        bold=True,
        spacing_after=8,
    )

    if TAARYA_LOGO.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(TAARYA_LOGO), width=Inches(2.2))

    add_centered(doc, "(ES-452: Major Project - Dissertation)", size=18, bold=True, spacing_after=16)
    add_centered(doc, "submitted in partial fulfillment of the requirement", size=16)
    add_centered(doc, "for the award of the degree of", size=16)
    add_centered(doc, "Bachelor of Technology", size=20, bold=True)
    add_centered(doc, "in", size=16)
    add_centered(doc, "Computer Science & Engineering", size=20, bold=True, spacing_after=12)
    add_centered(doc, "Submitted by", size=16, bold=True)
    add_centered(doc, "Amit Kumar", size=16, bold=True)
    add_centered(doc, "Enrollment No: 14618002722", size=16, spacing_after=12)
    add_centered(doc, "Under the supervision of", size=16, bold=True)
    add_centered(doc, "Dr. Madhumita Mahapatra", size=16, bold=True)
    add_centered(doc, "Professor, Department of Computer Science & Engineering", size=16, spacing_after=12)
    add_centered(doc, DEPARTMENT, size=18, bold=True)
    add_centered(doc, "Delhi Technical Campus", size=18, bold=True)
    add_centered(doc, "28/1, Knowledge Park-III, Greater Noida - 201306 (U.P.)", size=14)
    add_centered(doc, "April 2026", size=16, bold=True)


def add_front_matter(doc):
    doc.add_paragraph().add_run().add_break()
    doc.add_heading("DECLARATION", level=1)
    add_body(
        doc,
        "This is to certify that the material embodied in this Major Project - Dissertation titled "
        '"TaarYa: An Intelligent Agentic RAG-Driven Architecture for Astronomical Star Catalogs" '
        "being submitted in partial fulfillment of the requirements for the award of the degree of "
        "Bachelor of Technology in Computer Science & Engineering is based on my original work. "
        "It is further certified that this Major Project - Dissertation has not been submitted in full "
        "or in part to this university or any other university for the award of any other degree or diploma. "
        "All sources used in preparing this work have been duly acknowledged at the relevant places."
    )
    doc.add_paragraph()
    add_centered(doc, "(Amit Kumar)", size=12)
    add_centered(doc, "Enrollment No: 14618002722", size=12)
    doc.add_page_break()

    doc.add_heading("CERTIFICATE", level=1)
    add_body(
        doc,
        "This is to certify that the work embodied in this Major Project - Dissertation titled "
        '"TaarYa: An Intelligent Agentic RAG-Driven Architecture for Astronomical Star Catalogs" '
        "has been carried out by Amit Kumar (Enrollment No: 14618002722) under my supervision and guidance. "
        "To the best of my knowledge, the work reported in this dissertation is original, and has not been "
        "submitted elsewhere for the award of any degree or diploma."
    )
    doc.add_paragraph()
    add_centered(doc, "(Dr. Madhumita Mahapatra)", size=12)
    add_centered(doc, "Professor", size=12)
    add_centered(doc, DEPARTMENT, size=12)
    doc.add_paragraph()
    add_centered(doc, "Project Coordinator                                Head of the Department", size=12)
    add_centered(doc, "Department of CSE                                 Department of CSE", size=12)
    add_centered(doc, "Delhi Technical Campus, Greater Noida             Delhi Technical Campus, Greater Noida", size=12)
    doc.add_page_break()

    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    add_body(
        doc,
        "I express my sincere gratitude to Dr. Madhumita Mahapatra for her guidance, encouragement, "
        "and continuous support throughout the development of this major project. Her feedback helped shape "
        "both the technical direction and the final presentation of this work."
    )
    add_body(
        doc,
        "I also thank the faculty members of the Department of Computer Science & Engineering, Delhi Technical Campus, "
        "for providing the academic environment and resources required for this project. I am grateful to the project "
        "coordinator, the Head of the Department, and the institute administration for their support."
    )
    add_body(
        doc,
        "Finally, I acknowledge my family, friends, and peers for their constant motivation and encouragement during "
        "the completion of this project."
    )
    doc.add_page_break()

    doc.add_heading("ABSTRACT", level=1)
    add_body(
        doc,
        "TaarYa is a development-based major project that implements an intelligent, agentic Retrieval-Augmented "
        "Generation (RAG) architecture for astronomical star catalogs and scientific literature. The system is designed "
        "to address the practical difficulty faced by astronomy students and researchers who must otherwise switch between "
        "catalog databases, paper repositories, and programming tools in separate workflows."
    )
    add_body(
        doc,
        "The implemented system combines a FastAPI backend, a browser-based user interface, PostgreSQL with Q3C for "
        "spatial search, Qdrant for semantic paper retrieval, and Neo4j for graph-based relationships between stars, "
        "clusters, and papers. A language-model-driven agent routes user questions to the appropriate tools, such as "
        "cone search, semantic search, star lookup, and graph traversal, before synthesizing a grounded response."
    )
    add_body(
        doc,
        "TaarYa also includes ingestion pipelines for Gaia data, ArXiv papers, and generic catalog uploads, allowing "
        "the project to evolve from a fixed demo into an extensible astronomy research workspace. The frontend provides "
        "chat, exploration, system analysis, and research support interfaces, while the backend exposes modular APIs for "
        "search, ingestion, sessions, and agent interaction."
    )
    add_body(
        doc,
        "The report presents the analysis, design, implementation, and testing strategy of the system in the format "
        "prescribed for ES-452. Diagram spaces have been intentionally left in the report so that the final use case, "
        "data flow, class, ER, and activity diagrams may be inserted later without disturbing the report structure."
    )
    doc.add_page_break()

    doc.add_heading("TABLE OF CONTENTS", level=1)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Word field inserted below. Update fields after final edits if required.")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.italic = True
    toc_p = doc.add_paragraph()
    add_field(toc_p, r'TOC \o "1-3" \h \z \u', "Table of Contents")
    doc.add_page_break()

    doc.add_heading("LIST OF FIGURES", level=1)
    add_body(
        doc,
        "This section is reserved for the final list of figures. Update it after inserting the diagrams and screenshots "
        "into the report and finalizing their captions."
    )
    doc.add_page_break()

    doc.add_heading("LIST OF TABLES", level=1)
    add_body(
        doc,
        "This section is reserved for the final list of tables. Update it after finalizing the tabular content and captions."
    )


def add_chapter_1(doc):
    doc.add_heading("CHAPTER 1: Introduction", level=1)
    add_body(
        doc,
        "The scale of modern astronomy has transformed the discipline into a data-intensive science. Large sky surveys "
        "such as Gaia DR3 produce enormous quantities of astrometric and photometric observations, while scientific knowledge "
        "about those observations remains distributed across research papers, catalogs, and domain tools. Although the data is "
        "publicly available, the workflow required to use it effectively is fragmented."
    )
    add_body(
        doc,
        "A student or researcher typically has to query positional data through ADQL or SQL-like systems, interpret star metadata "
        "through catalog documentation, search literature separately through repositories such as ArXiv, and then use programming "
        "libraries such as Astropy or Pandas for analysis. This process is time-consuming and creates a barrier for users who do not "
        "already have strong database and scripting skills."
    )
    add_body(
        doc,
        "TaarYa addresses this gap by implementing an agentic RAG-based astronomy assistant that unifies multiple knowledge sources "
        "inside one application. Instead of treating astronomy archives as isolated repositories, the project treats them as an "
        "interactive environment through which an AI system can navigate, retrieve, cross-reference, and synthesize grounded answers."
    )
    add_body(
        doc,
        "The project combines four main ideas: explicit catalog ingestion, hybrid retrieval across spatial, semantic, and graph "
        "backends, an LLM-driven agent for tool selection, and a web-based interface for practical use. TaarYa therefore extends the "
        "proposal work from the previous semester into an implemented software system that supports real astronomical queries and "
        "research-oriented interaction."
    )


def add_chapter_2(doc):
    doc.add_heading("CHAPTER 2: Problem Statement", level=1)
    doc.add_heading("2.1 Problem Definition", level=2)
    add_body(
        doc,
        "The primary problem addressed by TaarYa is the fragmented workflow of astronomy data exploration. Astronomical star catalogs, "
        "paper repositories, and graph-like relationships between observations are usually accessed through separate tools. As a result, "
        "a single research question often requires multiple manual steps across different interfaces."
    )
    add_body(
        doc,
        "This fragmentation creates four practical difficulties. First, there is no unified interface for asking complex questions that "
        "combine position, catalog metadata, and literature context. Second, command-based systems such as ADQL and raw SQL present a "
        "technical barrier to students and early-career researchers. Third, semantic retrieval alone is not sufficient for astronomy "
        "because positional queries require strict geometric precision. Fourth, conventional search interfaces do not support multi-hop "
        "reasoning across stars, papers, and related entities."
    )
    add_body(
        doc,
        "A modern astronomy assistant must therefore support both precise structured retrieval and flexible natural language interaction. "
        "It must also expose the results in an understandable way and protect users from unsupported queries outside the available catalog coverage."
    )

    doc.add_heading("2.2 Objectives", level=2)
    objectives = [
        "To build a modular astronomy research assistant that integrates catalog data, scientific literature, and graph relationships within one system.",
        "To implement ingestion pipelines for Gaia data, ArXiv papers, and additional catalog uploads so that the platform remains extensible.",
        "To support hybrid retrieval by combining Q3C-powered spatial search, vector-based semantic search, and graph traversal.",
        "To develop an agent layer that can interpret natural language queries, choose appropriate tools, and synthesize grounded responses.",
        "To provide a browser-based dashboard through which users can chat with the system, explore stars, inspect system status, and navigate research views.",
        "To maintain a development structure that can be tested through ingestion, retrieval, API, and streaming-contract test scripts.",
    ]
    for item in objectives:
        add_bullet(doc, item)


def add_chapter_3(doc):
    doc.add_heading("CHAPTER 3: Analysis", level=1)
    doc.add_heading("3.1 Software Requirement Specifications", level=2)
    doc.add_heading("3.1.1 Functional Requirements of the Project", level=3)
    add_caption(doc, "Table 1: Functional Requirements")
    add_table(
        doc,
        ["Code", "Requirement"],
        [
            ("FR-01", "The system shall accept natural language astronomy queries through a web interface."),
            ("FR-02", "The system shall perform cone searches on stellar coordinates using PostgreSQL with Q3C."),
            ("FR-03", "The system shall support star lookup, nearby star search, and region-based counting."),
            ("FR-04", "The system shall ingest ArXiv papers into a vector database for semantic retrieval."),
            ("FR-05", "The system shall maintain graph relationships among stars, papers, and clusters through Neo4j."),
            ("FR-06", "The system shall expose APIs for hybrid search, ingestion, sessions, stars, papers, and the AI agent."),
            ("FR-07", "The system shall support background ingestion jobs for Gaia, ArXiv, and generic astronomy catalogs."),
            ("FR-08", "The system shall preserve session history for repeated user interaction."),
            ("FR-09", "The agent shall use real tool outputs instead of hallucinating unsupported astronomy results."),
            ("FR-10", "The system shall provide a dashboard with chat, exploration, and system analysis views."),
        ],
    )

    doc.add_heading("3.1.2 Non-functional Requirements of the Project", level=3)
    add_caption(doc, "Table 2: Non-functional Requirements")
    add_table(
        doc,
        ["Code", "Requirement"],
        [
            ("NFR-01", "The system should respond to common search requests within a practical interactive time for academic use."),
            ("NFR-02", "The architecture should be modular enough to let ingestion, retrieval, and agent components evolve independently."),
            ("NFR-03", "The interface should remain usable for students who are not experts in ADQL or database internals."),
            ("NFR-04", "The system should remain locally deployable with Dockerized services and optional local LLM backends."),
            ("NFR-05", "The project should be testable through lightweight local tests and service-backed validation scripts."),
            ("NFR-06", "The system should clearly indicate data coverage and avoid answering with unsupported catalog assumptions."),
        ],
    )

    doc.add_heading("3.2 Feasibility Study of the Project", level=2)
    add_body(
        doc,
        "Technical feasibility is high because the project uses mature open-source components for each subsystem: FastAPI for service orchestration, "
        "PostgreSQL plus Q3C for spatial indexing, Qdrant for vector similarity, Neo4j for graph relationships, and LangChain-compatible LLM layers for "
        "tool-driven orchestration. The current repository already implements these modules in a working codebase."
    )
    add_body(
        doc,
        "Operational feasibility is also strong because TaarYa is designed as a browser-accessible local application. The backend serves static pages and "
        "API routes from the same project, while ingestion jobs are triggered explicitly through API endpoints. This reduces accidental heavy processing on startup."
    )
    add_body(
        doc,
        "Economic feasibility is favorable because the system relies largely on open-source software, public astronomy datasets, and optional local LLM hosting. "
        "The project is suitable for academic environments where cost, reproducibility, and self-hosted data access are important."
    )
    add_body(
        doc,
        "Schedule feasibility remains realistic because the project is already in implementation form. The present task is report consolidation and final diagram insertion, "
        "not initial architecture design from scratch."
    )

    doc.add_heading("3.3 Tools / Technologies / Platform used", level=2)
    add_caption(doc, "Table 3: Technology Stack Used in TaarYa")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ("Backend Framework", "Python, FastAPI, Uvicorn"),
            ("Agent Layer", "LangChain-compatible tool orchestration with OpenAI or Ollama backends"),
            ("Spatial Database", "PostgreSQL 15 with Q3C"),
            ("Vector Store", "Qdrant"),
            ("Graph Database", "Neo4j"),
            ("Frontend", "HTML, CSS, JavaScript, browser-based dashboard pages"),
            ("Data / Analysis Libraries", "Astropy, Pandas, NumPy, sentence-transformers, arxiv"),
            ("Deployment Support", "Docker Compose, local virtual environment scripts"),
            ("Testing", "pytest and repository-level validation scripts"),
        ],
    )

    doc.add_heading("3.4 Use Case Diagrams / Data Flow Diagrams", level=2)
    add_body(
        doc,
        "The final report requires visual diagrams for the use case view and data flow views. The conceptual interactions are already clear from the implementation: "
        "the user interacts with the dashboard, the dashboard calls backend APIs, the backend routes requests to the retrieval and agent modules, and the retrieval "
        "modules interact with PostgreSQL, Qdrant, and Neo4j. The exact diagrams can be inserted later in the reserved spaces below."
    )
    add_caption(doc, "Figure 1: Use Case Diagram of TaarYa (placeholder)")
    add_placeholder(doc, "the Use Case Diagram", height_inches=2.8)
    add_caption(doc, "Figure 2: Data Flow Diagram Level 0 (placeholder)")
    add_placeholder(doc, "the Data Flow Diagram Level 0", height_inches=2.8)
    add_caption(doc, "Figure 3: Data Flow Diagram Level 1 (placeholder)")
    add_placeholder(doc, "the Data Flow Diagram Level 1", height_inches=2.8)


def add_chapter_4(doc):
    doc.add_heading("CHAPTER 4: Design and Architecture", level=1)
    doc.add_heading("4.1 Structure Chart / Work Breakdown Structure", level=2)
    add_body(
        doc,
        "The development of TaarYa can be decomposed into six major work streams: project setup and environment configuration, ingestion pipeline development, "
        "retrieval engine implementation, AI agent orchestration, web interface construction, and testing and validation."
    )
    add_bullet(doc, "Work Stream 1: configure environment, settings, Docker services, and application entry points.")
    add_bullet(doc, "Work Stream 2: implement Gaia, ArXiv, and generic catalog ingestion pipelines.")
    add_bullet(doc, "Work Stream 3: implement spatial, vector, graph, and hybrid retrieval modules.")
    add_bullet(doc, "Work Stream 4: implement the agent layer, tool wrappers, session history, and response synthesis.")
    add_bullet(doc, "Work Stream 5: build dashboard, chat, explore, analysis, and settings pages.")
    add_bullet(doc, "Work Stream 6: create API, ingestion, retrieval, backend, and streaming-contract test scripts.")
    add_caption(doc, "Figure 4: Work Breakdown Structure of TaarYa (placeholder)")
    add_placeholder(doc, "the Work Breakdown Structure", height_inches=2.6)

    doc.add_heading("4.2 Explanation of Modules", level=2)
    doc.add_heading("4.2.1 Application and API Layer", level=3)
    add_body(
        doc,
        "The main FastAPI application defined in src/main.py initializes database connections, mounts static files, and registers API routers. "
        "The API surface is organized by feature area, including stars, papers, search, sessions, ingestion, regions, and agent interaction."
    )
    doc.add_heading("4.2.2 Ingestion Module", level=3)
    add_body(
        doc,
        "The ingestion subsystem is responsible for bringing external knowledge into the system. Dedicated modules handle Gaia seeding, ArXiv paper ingestion, "
        "and generic catalog uploads. Background job tracking prevents duplicate ingestion requests while keeping the web interface responsive."
    )
    doc.add_heading("4.2.3 Retrieval Module", level=3)
    add_body(
        doc,
        "The retrieval layer is divided into spatial_search.py, vector_search.py, graph_search.py, and hybrid_search.py. Spatial search handles cone search, "
        "lookups, counts, and projection utilities. Vector search handles semantic similarity in Qdrant. Graph search manages star-paper-cluster relationships. "
        "Hybrid search acts as the unifying facade used by the API and agent layers."
    )
    doc.add_heading("4.2.4 Agent Module", level=3)
    add_body(
        doc,
        "The agent layer converts user questions into tool-driven workflows. It injects live coverage information into the system prompt, limits the number of tool "
        "calls, and instructs the model not to fabricate unsupported observations. This design makes the agent more dependable in a scientific setting."
    )
    doc.add_heading("4.2.5 Frontend Module", level=3)
    add_body(
        doc,
        "The static folder contains multiple user-facing pages including dashboard.html, chat.html, explore.html, analysis.html, literature_browser.html, "
        "research_chat.html, and settings.html. This indicates that TaarYa is designed as a complete user workspace rather than a single-purpose demo page."
    )
    doc.add_heading("4.2.6 Persistence and Session Module", level=3)
    add_body(
        doc,
        "The project maintains chat_sessions and chat_messages so that user conversations can persist across interactions. This supports research continuity and "
        "enables the agent to incorporate recent context into its answers."
    )

    doc.add_heading("4.3 Flow Chart / Activity Diagram", level=2)
    add_body(
        doc,
        "At runtime the workflow is straightforward: the user submits a query, the frontend sends it to the backend, the backend determines whether the request is "
        "a direct retrieval or an agent-driven interaction, the retrieval modules query the relevant backends, and the response is formatted for the interface. "
        "A detailed activity diagram should be inserted in the reserved space below."
    )
    add_caption(doc, "Figure 5: Activity Diagram / Query Flow of TaarYa (placeholder)")
    add_placeholder(doc, "the Activity Diagram", height_inches=2.8)

    doc.add_heading("4.4 ER Diagram / Class Diagram", level=2)
    add_body(
        doc,
        "The core persisted entities in the current implementation include stars, regions, chat sessions, and chat messages in the relational database; papers in the "
        "vector store; and stars, papers, and clusters in the graph database. A final ER diagram and class diagram can be inserted in the reserved spaces below."
    )
    add_caption(doc, "Figure 6: ER Diagram of TaarYa data entities (placeholder)")
    add_placeholder(doc, "the ER Diagram", height_inches=2.8)
    add_caption(doc, "Figure 7: Class Diagram of TaarYa modules (placeholder)")
    add_placeholder(doc, "the Class Diagram", height_inches=2.8)


def add_chapter_5(doc):
    doc.add_heading("CHAPTER 5: Implementation", level=1)
    add_body(
        doc,
        "This chapter presents the implemented software organization of TaarYa. The project is not limited to a concept note; it includes runnable source code, API modules, "
        "frontend pages, ingestion scripts, and tests. The implementation follows a layered architecture where each module focuses on a distinct responsibility."
    )

    doc.add_heading("5.1 Screenshots", level=2)
    add_body(
        doc,
        "The final version of the report should include screenshots of the dashboard, chat interface, exploration views, system status panels, and sample search results. "
        "These screenshots can be inserted later in the spaces below if required."
    )
    add_caption(doc, "Figure 8: TaarYa dashboard and chat interface (placeholder)")
    add_placeholder(doc, "the dashboard screenshot", height_inches=2.4)
    add_caption(doc, "Figure 9: TaarYa exploration and system status views (placeholder)")
    add_placeholder(doc, "the exploration screenshot", height_inches=2.4)

    doc.add_heading("5.2 Source Code of some modules", level=2)
    add_body(
        doc,
        "The following shortened excerpts summarize how the implemented modules are organized. They are adapted from the current source tree and included here to illustrate "
        "the development structure."
    )
    doc.add_heading("5.2.1 FastAPI application entry point", level=3)
    add_code_block(
        doc,
        """
app = FastAPI(title="TaarYa", version="0.2.0", lifespan=lifespan)
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
app.include_router(stars_router, prefix="/api")
app.include_router(papers_router, prefix="/api")
app.include_router(search_router, prefix="/api")
app.include_router(agent_router, prefix="/api")
app.include_router(ingestion_router, prefix="/api")
        """,
    )
    doc.add_heading("5.2.2 Background ingestion trigger", level=3)
    add_code_block(
        doc,
        """
@router.post("/gaia")
async def trigger_gaia_ingestion(background_tasks: BackgroundTasks):
    if _running["gaia"]:
        raise HTTPException(status_code=409, detail="Gaia ingestion is already running")
    background_tasks.add_task(loop.run_in_executor, executor, _run_gaia)
    return {"status": "started", "pipeline": "gaia"}
        """,
    )
    doc.add_heading("5.2.3 Hybrid retrieval facade", level=3)
    add_code_block(
        doc,
        """
class HybridSearch:
    def __init__(self):
        self.spatial = SpatialSearch()
        self.vector = VectorSearch()
        self.graph = GraphSearch()

    def multi_search(self, query_text=None, ra=None, dec=None, radius_deg=None, source_id=None, limit=20):
        ...
        """,
    )
    doc.add_heading("5.2.4 Agent service wrapper", level=3)
    add_code_block(
        doc,
        """
class AgentService:
    def ask(self, query, chat_history=None):
        from src.agent.agent import ask
        result = ask(query, chat_history)
        return result
        """,
    )


def add_chapter_6(doc):
    doc.add_heading("CHAPTER 6: Testing", level=1)
    doc.add_heading("6.1 Test Cases", level=2)
    add_body(
        doc,
        "The repository already contains a set of tests and validation scripts. Some are lightweight local tests that can run without all services, while others are "
        "service-backed validation scripts that assume a running server or active databases."
    )
    add_caption(doc, "Table 4: Planned and Implemented Test Coverage")
    add_table(
        doc,
        ["Test Area", "Purpose", "Source"],
        [
            ("Streaming contract", "Verify that scratchpad events are hidden and clean final answers are preserved.", "tests/test_streaming_contract.py"),
            ("Gaia parser", "Verify parsing of sample Gaia CSV data.", "tests/test_ingestion.py::test_gaia_parser"),
            ("API routes", "Exercise health, search, stars, papers, and stats endpoints against a running server.", "tests/test_api.py"),
            ("Retrieval layer", "Validate spatial, vector, and graph retrieval flows.", "tests/test_retrieval.py"),
            ("Backend smoke tests", "Quick checks for Qdrant and Neo4j.", "tests/test_backends.py"),
            ("API retry checks", "Re-run longer semantic and hybrid queries.", "tests/test_api_retry.py"),
        ],
    )

    doc.add_heading("6.2 Test Results", level=2)
    add_body(
        doc,
        "Two lightweight local validations were executed successfully inside the project virtual environment while preparing this report. "
        "The streaming-contract suite passed, confirming that the agent does not leak internal scratchpad events to the frontend and that its coverage prompt injection "
        "behavior remains testable. The Gaia parser unit test also passed, confirming successful parsing of sample catalog CSV data."
    )
    add_caption(doc, "Table 5: Local Validation Results Used for this Report")
    add_table(
        doc,
        ["Validation Command", "Observed Result", "Remarks"],
        [
            (r".\.venv\Scripts\python.exe -m pytest tests\test_streaming_contract.py -q", "5 passed", "Validated streaming behavior and prompt coverage logic."),
            (r".\.venv\Scripts\python.exe -m pytest tests\test_ingestion.py::test_gaia_parser -q", "1 passed", "Validated local CSV parsing pipeline for Gaia data."),
            ("Service-backed API and retrieval tests", "Available in repository", "Require configured services and, in some cases, a running local server."),
        ],
    )
    add_body(
        doc,
        "The presence of these tests, together with the modular source structure, indicates that the project has moved beyond conceptual design into a maintainable software "
        "implementation. Final end-to-end validation for submission should additionally include screenshots of successful live API interactions and system status pages."
    )


def add_chapter_7(doc):
    doc.add_heading("CHAPTER 7: Summary and Conclusion", level=1)
    add_body(
        doc,
        "TaarYa demonstrates how an astronomy-focused RAG system can be developed as a practical software product rather than only a theoretical proposal. "
        "The implemented architecture integrates structured astronomical data, semantic paper retrieval, graph context, and agentic interaction into one platform."
    )
    add_body(
        doc,
        "The project addresses a genuine academic need: enabling students and researchers to ask natural-language astronomy questions without losing the precision of "
        "catalog-based search. The codebase shows that this has been achieved through a modular implementation built on FastAPI, PostgreSQL with Q3C, Qdrant, Neo4j, and "
        "an LLM-driven orchestration layer."
    )
    add_body(
        doc,
        "From the perspective of a development-based major project, the most important outcome is the transition from concept to deployable system. TaarYa now contains "
        "a frontend workspace, background ingestion jobs, multi-backend retrieval services, agent routing, session support, and an evolving test suite."
    )


def add_chapter_8(doc):
    doc.add_heading("CHAPTER 8: Limitation of the Project and Future Work", level=1)
    doc.add_heading("8.1 Limitation", level=2)
    limitations = [
        "The project depends on multiple backing services such as PostgreSQL, Qdrant, Neo4j, and an LLM runtime, which increases deployment complexity.",
        "Catalog coverage is limited by the data currently ingested into the local database and does not automatically span the full sky.",
        "Some richer astronomy workflows still depend on user interpretation of the returned data rather than full autonomous analysis pipelines.",
        "The final research diagrams and polished screenshot set still need to be inserted into the report.",
    ]
    for item in limitations:
        add_bullet(doc, item)

    doc.add_heading("8.2 Future Work", level=2)
    future_work = [
        "Expand ingestion beyond Gaia and ArXiv to include additional astronomy catalogs and better cross-catalog normalization.",
        "Strengthen the graph layer with more explicit paper-to-paper citation and star-to-region relationships.",
        "Add richer code-execution and plotting workflows so that TaarYa can generate more advanced astronomy visualizations on demand.",
        "Improve frontend workflows for research history, result export, and guided exploration of discovered stars and papers.",
        "Complete full end-to-end live validation across all backends and document benchmark-oriented performance measurements.",
    ]
    for item in future_work:
        add_bullet(doc, item)


def add_bibliography(doc):
    doc.add_heading("BIBLIOGRAPHY", level=1)
    references = [
        '[1] Y. Gao et al., "Retrieval-Augmented Generation for Large Language Models: A Comprehensive Review," arXiv:2312.10997, 2024.',
        '[2] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP," NeurIPS, vol. 33, pp. 9459-9474, 2020.',
        '[3] G. Edelman et al., "GraphRAG: Towards Graph-Based Retrieval-Augmented Generation," arXiv:2404.16130, 2024.',
        '[4] Gaia Collaboration, "Gaia Data Release 3 Documentation."',
        '[5] S. Koposov et al., "Q3C, a Quad Tree Cube - an Astronomical Data Structure for Massive Sky Surveys," Astronomy and Computing, 2010.',
        '[6] A. Vaswani et al., "Attention Is All You Need," NeurIPS, 2017.',
        '[7] T. Brown et al., "Language Models are Few-Shot Learners," NeurIPS, 2020.',
        '[8] TaarYa Project Repository, README and implementation modules, 2026.',
        '[9] TaarYa Project Synopsis, "Developing an Intelligent RAG-Driven Architecture for Astronomical Star Catalogs," Delhi Technical Campus, 2026.',
    ]
    for item in references:
        add_body(doc, item)


def add_appendix(doc):
    doc.add_heading("APPENDIX", level=1)
    add_body(
        doc,
        "Appendix A contains a concise list of key API routes and implementation artifacts that support the current version of TaarYa."
    )
    add_caption(doc, "Table 6: Key API Endpoints in TaarYa")
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ("GET /health", "Basic backend health check."),
            ("GET /api/stars/cone-search", "Retrieve stars around an RA/Dec center."),
            ("GET /api/stars/lookup/{source_id}", "Retrieve a specific star by source ID."),
            ("GET /api/papers/search", "Semantic search of ingested papers."),
            ("GET /api/search/hybrid", "Combined spatial / semantic / graph retrieval."),
            ("GET /api/stats", "System statistics for PostgreSQL, Qdrant, and Neo4j."),
            ("POST /api/ingest/gaia", "Trigger Gaia ingestion in the background."),
            ("POST /api/ingest/arxiv", "Trigger ArXiv ingestion in the background."),
            ("POST /api/ingest/catalog/upload", "Upload and ingest a generic astronomy catalog file."),
        ],
    )
    add_body(
        doc,
        "Appendix B is reserved for the plagiarism report. As per the written institute rules, the similarity percentage should remain below 20 percent."
    )
    doc.add_page_break()
    doc.add_heading("PLAGIARISM REPORT PLACEHOLDER", level=1)
    add_placeholder(doc, "the plagiarism report page(s)", height_inches=4.0)


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    add_cover_page(doc)

    prelim_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(prelim_section)
    configure_footer(prelim_section, fmt="roman", start=1)
    add_front_matter(doc)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(body_section)
    configure_footer(body_section, fmt="decimal", start=1)

    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_chapter_7(doc)
    add_chapter_8(doc)
    add_bibliography(doc)
    add_appendix(doc)

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_report()
